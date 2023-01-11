import docx
from docx.shared import Cm, Pt

document = docx.Document()

#_______1
section = document.sections[-1]
section.top_margin = Cm(2) #Верхний отступ
section.bottom_margin = Cm(2) #Нижний отступ
section.left_margin = Cm(3) #Отступ слева
section.right_margin = Cm(2) #Отступ справа
#_______2
paragraph_format = document.styles['Normal'].paragraph_format
paragraph_format.line_spacing = Pt(12) #межстрочный интервал
#_______3
style = document.styles['Normal']
font = style.font
font.name ='Times New Roman' #Стиль шрифта
font.size = Pt(12) #Размер шрифта

document.add_paragraph('Приветики!')

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)


records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.save('/home/lines14/projects/judicial_telegram_bot/documents/example.docx')