# document 1:
# petition for familiarization with the case materials

from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
import template_reserve

def make_table_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def make_table_columns_align_right(*columns):
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_table_rows_height_0_5(*rows):
    for row in rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(0.5)

# font = paragraph.runs[0].font
# font.size= Pt(10)
# style = document.styles['Normal']
# style.paragraph_format.line_spacing = Pt(8)

document = Document()

section = document.sections[0]
section.top_margin = Cm(2) #Верхний отступ
section.bottom_margin = Cm(2) #Нижний отступ
section.left_margin = Cm(3) #Отступ слева
section.right_margin = Cm(2) #Отступ справа

# paragraph_format = document.styles['Normal'].paragraph_format
# paragraph_format.line_spacing = Pt(12) #межстрочный интервал

# style = document.styles['Normal']
# font = style.font
# font.name ='Times New Roman' #Стиль шрифта
# font.size = Pt(12) #Размер шрифта

records = (
    (template_reserve.VAR1, template_reserve.VAR2),
    (template_reserve.VAR3, template_reserve.VAR4),
    (template_reserve.VAR5, template_reserve.VAR6),
    (template_reserve.VAR7, template_reserve.VAR8),
    (template_reserve.VAR9, template_reserve.VAR10),
    (template_reserve.VAR11, template_reserve.VAR12),
    (template_reserve.VAR13, template_reserve.VAR14),
    (template_reserve.VAR15, template_reserve.VAR16),
    (template_reserve.VAR17, template_reserve.VAR18),
    (template_reserve.VAR19, template_reserve.VAR20)
)

table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = template_reserve.VAR21
hdr_cells[1].text = template_reserve.VAR22
for first_col, second_col in records:
    row_cells = table.add_row().cells
    row_cells[0].text = first_col
    row_cells[1].text = second_col

make_table_rows_bold(table.rows[0], table.rows[2], table.rows[7])
make_table_columns_align_right(table.columns[0])
set_table_rows_height_0_5(table.rows[9])

para = document.add_paragraph(template_reserve.VAR23)
para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
para.paragraph_format.space_after = Pt(8)

para = document.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_after = Pt(8)
para.add_run(template_reserve.VAR24).bold = True


para = document.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_after = Pt(8)
para.add_run(template_reserve.VAR25).bold = True


para = document.add_paragraph(f'\t{template_reserve.VAR26}')
para.paragraph_format.space_after = Pt(8)

para = document.add_paragraph(f'\t{template_reserve.VAR27}')
para.paragraph_format.space_after = Pt(8)

para = document.add_paragraph(f'\t{template_reserve.VAR28}')
para.paragraph_format.space_after = Pt(8)

para = document.add_paragraph(template_reserve.VAR29)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_after = Pt(8)

para = document.add_paragraph(f'\t{template_reserve.VAR30}\n')
para.paragraph_format.space_after = Pt(8)

table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = template_reserve.VAR31
hdr_cells[1].text = template_reserve.VAR32

make_table_rows_bold(table.rows[0])
make_table_columns_align_right(table.columns[1])

document.save('/home/lines14/projects/judicial_telegram_bot/documents/judicial_writer_1.docx')