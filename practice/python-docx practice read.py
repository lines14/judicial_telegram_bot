import docx
from pathlib import Path
from modules import read_all_file
destination = Path(__file__).resolve().parent.parent

# doc = docx.Document('/home/lines14/projects/judicial_telegram_bot/judicial_№1_writer.docx')

# print(len(doc.paragraphs))

# print(doc.paragraphs[0].text)

# print(len(doc.paragraphs[0].runs))

# print(doc.paragraphs[0].runs[0].text)
# print(doc.paragraphs[0].runs[1].text)
# print(doc.paragraphs[0].runs[2].text)
# print(doc.paragraphs[0].runs[3].text)

print(read_all_file.gettext(f'{destination}/documents/judicial_writer_1.docx'))