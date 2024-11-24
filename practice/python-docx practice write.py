import docx
from pathlib import Path
path = Path(__file__).resolve().parent.parent

doc = docx.Document()

p = doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

doc.save(f'{path}/documents/practice.docx')