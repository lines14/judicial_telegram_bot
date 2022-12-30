# petition for familiarization with the case materials:

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import dotenv

dotenv.load_dotenv()

VAR1 = os.environ.get('VAR1')
VAR2 = os.environ.get('VAR2')
VAR3 = os.environ.get('VAR3')
VAR4 = os.environ.get('VAR4')
VAR5 = os.environ.get('VAR5')
VAR6 = os.environ.get('VAR6')
VAR7 = os.environ.get('VAR7')
VAR8 = os.environ.get('VAR8')
VAR9 = os.environ.get('VAR9')
VAR10 = os.environ.get('VAR10')
VAR11 = os.environ.get('VAR11')
VAR12 = os.environ.get('VAR12')
VAR13 = os.environ.get('VAR13')
VAR14 = os.environ.get('VAR14')
VAR15 = os.environ.get('VAR15')
VAR16 = os.environ.get('VAR16')
VAR17 = os.environ.get('VAR17')
VAR18 = os.environ.get('VAR18')
VAR19 = os.environ.get('VAR19')
VAR20 = os.environ.get('VAR20')
VAR21 = os.environ.get('VAR21')
VAR22 = os.environ.get('VAR22')
VAR23 = os.environ.get('VAR23')
VAR24 = os.environ.get('VAR24')
VAR25 = os.environ.get('VAR25')
VAR26 = os.environ.get('VAR26')
VAR27 = os.environ.get('VAR27')
VAR28 = os.environ.get('VAR28')
VAR29 = os.environ.get('VAR29')
VAR30 = os.environ.get('VAR30')
VAR31 = os.environ.get('VAR31')
VAR32 = os.environ.get('VAR32')

def make_table_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

#     font = paragraph.runs[0].font
#     font.size= Pt(10)

def make_table_columns_align_right(*columns):
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

document = docx.Document()

records = (
    (VAR1, VAR2),
    (VAR3, VAR4),
    (VAR5, VAR6),
    (VAR7, VAR8),
    (VAR9, VAR10),
    (VAR11, VAR12),
    (VAR13, VAR14),
    (VAR15, VAR16),
    (VAR17, VAR18),
    (VAR19, VAR20)
)

table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = VAR21
hdr_cells[1].text = VAR22
for first_col, second_col in records:
    row_cells = table.add_row().cells
    row_cells[0].text = first_col
    row_cells[1].text = second_col
make_table_rows_bold(table.rows[0], table.rows[2], table.rows[7])
make_table_columns_align_right(table.columns[0])

document.add_paragraph(VAR23)
para = document.add_paragraph('')
para.add_run(VAR24).bold = True
para = document.add_paragraph('')
para.add_run(VAR25).bold = True
document.add_paragraph(VAR26)
document.add_paragraph(VAR27)
document.add_paragraph(VAR28)
document.add_paragraph(VAR29)
document.add_paragraph(VAR30)

table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = VAR31
hdr_cells[1].text = VAR32
make_table_rows_bold(table.rows[0])
make_table_columns_align_right(table.columns[1])

document.save('/home/lines14/projects/judicial_telegram_bot/judicial_№1_writer.docx')