import docx
from modules import read_all_file

doc = docx.Document('/home/lines14/projects/judicial_telegram_bot/practice.docx')

print(len(doc.paragraphs))

print(doc.paragraphs[0].text)

print(len(doc.paragraphs[0].runs))

print(doc.paragraphs[0].runs[0].text)
print(doc.paragraphs[0].runs[1].text)
print(doc.paragraphs[0].runs[2].text)
print(doc.paragraphs[0].runs[3].text)

print(read_all_file.gettext('/home/lines14/projects/judicial_telegram_bot/judicial_template.docx'))